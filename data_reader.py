import warnings
import re
from typing import Generator
import zipfile
from pathlib import Path
from dataclasses import dataclass

import docx
import spacy
import nltk
import tqdm
import pandas as pd
from nltk.corpus import stopwords


@dataclass
class Document:
    name: str
    text: list[str]


class DocxDataReader:
    def __init__(self, zipfile_path: Path):
        self.path = zipfile_path

        self.nlp = spacy.load('ru_core_news_md')

        nltk.download('stopwords')
        self.stop_words = stopwords.words('russian')
        self.stop_words.extend(['\n', ',', '-', '.', '"', ')', '(', ';', '—', '/', '\t', ':'])

    def _read_docx_from_zipfile(self):
        archive = zipfile.ZipFile(self.path, 'r')

        for name in archive.namelist():
            if not re.fullmatch(r'.*\.docx', name):
                continue

            with archive.open(name) as file:
                doc = docx.Document(file)
                contents = {'paragraphs': [p.text for p in doc.paragraphs if len(p.text) > 0]}
            yield name, pd.DataFrame(data=contents)

    def _enrich_frame(self, df):
        result = df.copy()
        tokens = []
        lemma = []
        pos = []
        col_to_parse = 'paragraphs'

        for x in self.nlp.pipe(df[col_to_parse].astype('unicode').values, batch_size=50):
            if x.is_parsed:
                tokens.append([n.text for n in x])
                lemma.append([n.lemma_ for n in x])
                pos.append([n.pos_ for n in x])
            else:
                tokens.append(None)
                lemma.append(None)
                pos.append(None)

        result['tokens'] = tokens
        result['lemma'] = lemma
        result['pos'] = pos

        return result

    def _prepare_for_LDA(self, text_df):
        with warnings.catch_warnings():
            warnings.simplefilter('ignore')

            enriched = self._enrich_frame(text_df)

        text = []
        for sentence in enriched.lemma:
            text.extend([word for word in sentence if word not in self.stop_words])
        return text

    def read_documents(self) -> Generator[Document, None, None]:
        doc_reader = self._read_docx_from_zipfile()
        for name, df in tqdm.tqdm(doc_reader):
            yield Document(name, self._prepare_for_LDA(df))
